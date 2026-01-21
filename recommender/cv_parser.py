"""
Module pour parser les CVs (PDF, DOCX) et extraire le texte
"""
import re
from pathlib import Path
from typing import Optional

import PyPDF2
import pdfplumber
from docx import Document


class CVParser:
    """Extracteur de texte depuis des fichiers CV (PDF, DOCX)"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_cv(self, file_path: str) -> str:
        """
        Parse un CV et retourne le texte extrait
        
        Args:
            file_path: Chemin vers le fichier CV
            
        Returns:
            Texte extrait du CV
            
        Raises:
            ValueError: Si le format n'est pas supporté
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif extension == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(
                f"Format non supporté: {extension}. "
                f"Formats acceptés: {', '.join(self.supported_formats)}"
            )
    
    def parse_cv_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Parse un CV depuis des bytes (pour upload web)
        
        Args:
            file_bytes: Contenu du fichier en bytes
            filename: Nom du fichier (pour déterminer l'extension)
            
        Returns:
            Texte extrait
        """
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf_bytes(file_bytes)
        elif extension in ['.docx', '.doc']:
            return self._parse_docx_bytes(file_bytes)
        elif extension == '.txt':
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Format non supporté: {extension}")
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse un fichier PDF avec pdfplumber (meilleur que PyPDF2)"""
        text = []
        
        try:
            # Essayer avec pdfplumber (meilleur pour les PDFs complexes)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            if text:
                return '\n'.join(text)
        except Exception as e:
            print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
        # Fallback vers PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Impossible de parser le PDF: {e}")
    
    def _parse_pdf_bytes(self, file_bytes: bytes) -> str:
        """Parse un PDF depuis des bytes"""
        from io import BytesIO
        
        text = []
        
        try:
            # Essayer avec pdfplumber
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            
            if text:
                return '\n'.join(text)
        except Exception:
            pass
        
        # Fallback vers PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Impossible de parser le PDF: {e}")
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse un fichier DOCX"""
        try:
            doc = Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Impossible de parser le DOCX: {e}")
    
    def _parse_docx_bytes(self, file_bytes: bytes) -> str:
        """Parse un DOCX depuis des bytes"""
        from io import BytesIO
        
        try:
            doc = Document(BytesIO(file_bytes))
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Impossible de parser le DOCX: {e}")
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse un fichier TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Essayer avec un encodage différent
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def clean_text(self, text: str) -> str:
        """
        Nettoie le texte extrait du CV
        
        Args:
            text: Texte brut
            
        Returns:
            Texte nettoyé
        """
        # Supprimer les caractères spéciaux excessifs
        text = re.sub(r'\s+', ' ', text)  # Multiple espaces -> un seul
        text = re.sub(r'\n+', '\n', text)  # Multiple newlines -> un seul
        
        # Supprimer les bullet points
        text = re.sub(r'[•◦▪▫●○■□▪▫]', '', text)
        
        # Nettoyer les espaces autour des ponctuations
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        
        return text.strip()


# Fonction utilitaire
def extract_text_from_cv(file_path: str) -> str:
    """
    Fonction helper pour extraire rapidement le texte d'un CV
    
    Args:
        file_path: Chemin vers le CV
        
    Returns:
        Texte nettoyé du CV
    """
    parser = CVParser()
    raw_text = parser.parse_cv(file_path)
    return parser.clean_text(raw_text)


if __name__ == "__main__":
    # Test du parser
    import sys
    
    if len(sys.argv) > 1:
        cv_path = sys.argv[1]
        parser = CVParser()
        
        try:
            text = parser.parse_cv(cv_path)
            cleaned_text = parser.clean_text(text)
            
            print("-" * 80)
            print(f"CV parse : {cv_path}")
            print("-" * 80)
            print(cleaned_text[:500])  # Afficher les 500 premiers caractères
            print(f"\n... (Total: {len(cleaned_text)} caractères)")
        except Exception as e:
            print(f"Erreur: {e}")
    else:
        print("Usage: python cv_parser.py <chemin_vers_cv>")
